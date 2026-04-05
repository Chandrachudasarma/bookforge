"""DOCX Exporter — Stage 6 for DOCX output.

Converts BookManifest to a print-ready Word document via python-docx.
Three mandatory post-processing steps:
  1. Build document from HTML sections
  2. Apply hairline grid borders to all tables
  3. Apply headers (book title) and footers (page numbers)
"""

from __future__ import annotations

from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from bookforge.core.exceptions import ExportError
from bookforge.core.models import (
    BookManifest,
    ExportResult,
    SectionRole,
    ValidationResult,
)
from bookforge.core.registry import register_exporter
from bookforge.export.base import BaseExporter
from bookforge.export.docx_table_borders import apply_table_borders


@register_exporter("docx")
class DocxExporter(BaseExporter):
    """Renders BookManifest → print-ready DOCX via python-docx."""

    output_format = "docx"

    def export(
        self,
        manifest: BookManifest,
        template=None,
        output_path: Path = None,
    ) -> ExportResult:
        if output_path is None:
            raise ExportError("output_path is required")

        try:
            return self._export(manifest, template, output_path)
        except ExportError:
            raise
        except Exception as exc:
            raise ExportError(f"DOCX export failed: {exc}") from exc

    def _export(self, manifest: BookManifest, template, output_path: Path) -> ExportResult:
        doc = Document()
        meta = manifest.metadata

        # --- Page margins ---
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.875)
            section.right_margin = Inches(0.625)

        # --- Render each section ---
        for book_section in sorted(manifest.sections, key=lambda s: s.order):
            self._render_section(doc, book_section)
            if book_section.role == SectionRole.CHAPTER:
                doc.add_page_break()

        # --- Apply table borders (mandatory) ---
        apply_table_borders(doc)

        # --- Apply headers and footers ---
        title = meta.title if meta else "Untitled"
        self._apply_headers_footers(doc, title)

        # --- Save ---
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return ExportResult(format="docx", output_path=output_path, success=True)

    def validate(self, output_path: Path) -> ValidationResult:
        """Basic validation: check the file can be opened by python-docx."""
        try:
            Document(str(output_path))
            return ValidationResult(valid=True)
        except Exception as exc:
            return ValidationResult(valid=False, errors=[str(exc)])

    # ------------------------------------------------------------------
    # Section rendering
    # ------------------------------------------------------------------

    def _render_section(self, doc: Document, section) -> None:
        """Render one BookSection into the Document."""
        soup = BeautifulSoup(section.content_html, "lxml")
        body = soup.find("body") or soup

        for element in body.find_all(True, recursive=False):
            self._render_element(doc, element)

    def _render_element(self, doc: Document, element) -> None:
        tag = element.name

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            style = f"Heading {level}"
            try:
                para = doc.add_paragraph(style=style)
            except KeyError:
                para = doc.add_paragraph()
            para.add_run(element.get_text(strip=True)).bold = (level == 1)

        elif tag == "p":
            text = element.get_text(separator="", strip=False)
            if text.strip():
                para = doc.add_paragraph()
                _add_inline_formatted_runs(para, element)

        elif tag in ("table",):
            self._render_table(doc, element)

        elif tag == "figure":
            caption = element.find("figcaption")
            img = element.find("img")
            if img:
                alt = img.get("alt", "")
                if alt:
                    doc.add_paragraph(f"[Image: {alt}]")
            if caption:
                cap_para = doc.add_paragraph(caption.get_text(strip=True))
                cap_para.style = "Caption" if "Caption" in [s.name for s in doc.styles] else "Normal"

        elif tag in ("ul", "ol"):
            for li in element.find_all("li", recursive=False):
                para = doc.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
                para.add_run(li.get_text(strip=True))

        elif tag in ("pre", "code"):
            doc.add_paragraph(element.get_text(strip=True), style="Normal")

        elif tag in ("section", "article", "div", "nav"):
            for child in element.find_all(True, recursive=False):
                self._render_element(doc, child)

        elif tag == "hr":
            doc.add_paragraph("─" * 40)

        elif tag == "blockquote":
            para = doc.add_paragraph(element.get_text(strip=True))
            para.paragraph_format.left_indent = Inches(0.5)

    def _render_table(self, doc: Document, table_el) -> None:
        """Convert an HTML table element to a python-docx table."""
        rows = table_el.find_all("tr")
        if not rows:
            return

        col_count = max(
            len(row.find_all(["td", "th"])) for row in rows
        )
        if col_count == 0:
            return

        table = doc.add_table(rows=0, cols=col_count)

        for row_el in rows:
            cells = row_el.find_all(["td", "th"])
            row = table.add_row()
            for i, cell_el in enumerate(cells[:col_count]):
                cell = row.cells[i]
                cell.text = cell_el.get_text(strip=True)
                if cell_el.name == "th":
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    # ------------------------------------------------------------------
    # Headers and footers
    # ------------------------------------------------------------------

    def _apply_headers_footers(self, doc: Document, title: str) -> None:
        """Add book title in header and page number in footer."""
        for section in doc.sections:
            # Header: book title
            header = section.header
            if not header.paragraphs:
                header.add_paragraph()
            header_para = header.paragraphs[0]
            header_para.text = title
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Footer: page number
            footer = section.footer
            if not footer.paragraphs:
                footer.add_paragraph()
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.clear()
            _add_page_number_field(footer_para)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _add_inline_formatted_runs(para: Paragraph, element) -> None:
    """Add runs with bold/italic formatting from inline HTML elements."""
    for child in element.children:
        if hasattr(child, "name"):
            text = child.get_text(separator="", strip=False)
            run = para.add_run(text)
            if child.name in ("strong", "b"):
                run.bold = True
            elif child.name in ("em", "i"):
                run.italic = True
        else:
            # NavigableString
            text = str(child)
            if text:
                para.add_run(text)


def _add_page_number_field(para: Paragraph) -> None:
    """Insert a PAGE field code into the paragraph for auto page numbers."""
    run = para.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
