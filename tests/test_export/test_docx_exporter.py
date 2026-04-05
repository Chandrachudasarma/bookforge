"""Tests for the DOCX exporter (Stage 6).

Verifies: section rendering, table borders, headers/footers, page numbers.
"""

from pathlib import Path

import pytest

docx_mod = pytest.importorskip("docx", reason="python-docx not installed")
from docx import Document
from docx.oxml.ns import qn

from bookforge.core.models import (
    Asset,
    BookManifest,
    BookMetadata,
    BookSection,
    ExportResult,
    SectionRole,
)
from bookforge.export.docx_exporter import DocxExporter


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def basic_manifest() -> BookManifest:
    return BookManifest(
        sections=[
            BookSection(
                role=SectionRole.CHAPTER,
                title="Introduction",
                content_html="<h1>Introduction</h1><p>First paragraph.</p>",
                order=1,
            ),
        ],
        metadata=BookMetadata(title="Test Book", authors=["Jane Doe"]),
        assets=[],
    )


@pytest.fixture
def manifest_with_table() -> BookManifest:
    table_html = (
        "<h1>Data</h1>"
        "<table>"
        "<thead><tr><th>Name</th><th>Value</th></tr></thead>"
        "<tbody>"
        "<tr><td>Alpha</td><td>1.0</td></tr>"
        "<tr><td>Beta</td><td>2.0</td></tr>"
        "</tbody>"
        "</table>"
    )
    return BookManifest(
        sections=[
            BookSection(
                role=SectionRole.CHAPTER,
                title="Data Chapter",
                content_html=table_html,
                order=1,
            ),
        ],
        metadata=BookMetadata(title="Table Book", authors=["John Smith"]),
        assets=[],
    )


@pytest.fixture
def multi_section_manifest() -> BookManifest:
    return BookManifest(
        sections=[
            BookSection(
                role=SectionRole.TITLE_PAGE,
                title="Title",
                content_html="<h1>My Book</h1><p>By Author</p>",
                order=0,
            ),
            BookSection(
                role=SectionRole.CHAPTER,
                title="Chapter 1",
                content_html="<h1>Chapter 1</h1><p>Content.</p><h2>Section A</h2><p>More content.</p>",
                order=1,
            ),
            BookSection(
                role=SectionRole.CHAPTER,
                title="Chapter 2",
                content_html="<h1>Chapter 2</h1><p>Second chapter.</p>",
                order=2,
            ),
        ],
        metadata=BookMetadata(title="Multi Chapter", authors=["Author"]),
        assets=[],
    )


# ---------------------------------------------------------------------------
# Tests — basic export
# ---------------------------------------------------------------------------


def test_exports_valid_docx(basic_manifest, tmp_path):
    exporter = DocxExporter()
    out = tmp_path / "output.docx"
    result = exporter.export(basic_manifest, template=None, output_path=out)

    assert result.success
    assert out.exists()
    assert out.stat().st_size > 0

    # Verify it opens with python-docx
    doc = Document(str(out))
    assert len(doc.paragraphs) > 0


def test_renders_heading_and_paragraph(basic_manifest, tmp_path):
    exporter = DocxExporter()
    out = tmp_path / "output.docx"
    exporter.export(basic_manifest, template=None, output_path=out)

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("Introduction" in t for t in texts)
    assert any("First paragraph" in t for t in texts)


# ---------------------------------------------------------------------------
# Tests — table borders
# ---------------------------------------------------------------------------


def test_table_has_borders(manifest_with_table, tmp_path):
    exporter = DocxExporter()
    out = tmp_path / "output.docx"
    exporter.export(manifest_with_table, template=None, output_path=out)

    doc = Document(str(out))
    assert len(doc.tables) >= 1

    table = doc.tables[0]
    # Check that at least one cell has border XML
    cell = table.rows[0].cells[0]
    tc = cell._tc
    borders = tc.findall(".//" + qn("w:tcBorders"))
    assert len(borders) > 0, "Table cell should have border definitions"


def test_table_content_correct(manifest_with_table, tmp_path):
    exporter = DocxExporter()
    out = tmp_path / "output.docx"
    exporter.export(manifest_with_table, template=None, output_path=out)

    doc = Document(str(out))
    table = doc.tables[0]

    # Check cell content
    assert table.rows[0].cells[0].text == "Name"
    assert table.rows[0].cells[1].text == "Value"
    assert table.rows[1].cells[0].text == "Alpha"
    assert table.rows[2].cells[0].text == "Beta"


# ---------------------------------------------------------------------------
# Tests — headers / footers
# ---------------------------------------------------------------------------


def test_header_contains_book_title(basic_manifest, tmp_path):
    exporter = DocxExporter()
    out = tmp_path / "output.docx"
    exporter.export(basic_manifest, template=None, output_path=out)

    doc = Document(str(out))
    for section in doc.sections:
        header_text = section.header.paragraphs[0].text
        assert "Test Book" in header_text


def test_footer_has_page_number_field(basic_manifest, tmp_path):
    exporter = DocxExporter()
    out = tmp_path / "output.docx"
    exporter.export(basic_manifest, template=None, output_path=out)

    doc = Document(str(out))
    for section in doc.sections:
        footer_xml = section.footer._element.xml
        # PAGE field code should be present
        assert "PAGE" in footer_xml


# ---------------------------------------------------------------------------
# Tests — multi-section ordering
# ---------------------------------------------------------------------------


def test_multi_section_renders_all_chapters(multi_section_manifest, tmp_path):
    exporter = DocxExporter()
    out = tmp_path / "output.docx"
    exporter.export(multi_section_manifest, template=None, output_path=out)

    doc = Document(str(out))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "Chapter 1" in all_text
    assert "Chapter 2" in all_text
    assert "My Book" in all_text


# ---------------------------------------------------------------------------
# Tests — validation
# ---------------------------------------------------------------------------


def test_validate_valid_docx(basic_manifest, tmp_path):
    exporter = DocxExporter()
    out = tmp_path / "output.docx"
    exporter.export(basic_manifest, template=None, output_path=out)

    result = exporter.validate(out)
    assert result.valid


def test_validate_invalid_file(tmp_path):
    exporter = DocxExporter()
    bad_file = tmp_path / "bad.docx"
    bad_file.write_text("not a real docx")

    result = exporter.validate(bad_file)
    assert not result.valid
