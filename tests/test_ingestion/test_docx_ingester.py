"""Tests for the DOCX ingester (Stage 1).

Creates real .docx files via python-docx for testing.
"""

from pathlib import Path

import pytest

docx_mod = pytest.importorskip("docx", reason="python-docx not installed")
from docx import Document
from docx.shared import Pt

from bookforge.ingestion.docx_ingester import DocxIngester


# ---------------------------------------------------------------------------
# Fixtures — create real DOCX files
# ---------------------------------------------------------------------------


@pytest.fixture
def simple_docx(tmp_path: Path) -> Path:
    """DOCX with Heading 1, paragraphs, bold/italic runs."""
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("First paragraph of the introduction.")
    p = doc.add_paragraph()
    p.add_run("Some ").bold = False
    run_b = p.add_run("bold text")
    run_b.bold = True
    p.add_run(" and ")
    run_i = p.add_run("italic text")
    run_i.italic = True
    p.add_run(".")
    doc.add_heading("Methods", level=2)
    doc.add_paragraph("Methods paragraph.")

    path = tmp_path / "simple.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_table(tmp_path: Path) -> Path:
    """DOCX with a table (header row + data rows)."""
    doc = Document()
    doc.add_heading("Data Table", level=1)

    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Alpha"
    table.rows[1].cells[1].text = "1.0"
    table.rows[2].cells[0].text = "Beta"
    table.rows[2].cells[1].text = "2.0"

    path = tmp_path / "with_table.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_config(tmp_path: Path) -> dict:
    return {
        "pipeline": {"temp_dir": str(tmp_path / "temp")},
    }


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


def test_ingests_simple_docx(simple_docx, docx_config):
    ingester = DocxIngester()
    raw = ingester.ingest(simple_docx, docx_config)

    assert raw.format_hint == "html"
    assert raw.source_path == simple_docx
    assert raw.source_metadata.get("original_format") == "docx"


def test_maps_heading_styles(simple_docx, docx_config):
    ingester = DocxIngester()
    raw = ingester.ingest(simple_docx, docx_config)

    assert "<h1>" in raw.text
    assert "Introduction" in raw.text
    assert "<h2>" in raw.text
    assert "Methods" in raw.text


def test_preserves_inline_formatting(simple_docx, docx_config):
    ingester = DocxIngester()
    raw = ingester.ingest(simple_docx, docx_config)

    assert "<strong>" in raw.text
    assert "bold text" in raw.text
    assert "<em>" in raw.text
    assert "italic text" in raw.text


def test_renders_table_as_html(docx_with_table, docx_config):
    ingester = DocxIngester()
    raw = ingester.ingest(docx_with_table, docx_config)

    assert "<table>" in raw.text
    assert "<th>" in raw.text  # first row rendered as header
    assert "Alpha" in raw.text
    assert "Beta" in raw.text


def test_can_handle_returns_true_for_docx():
    ingester = DocxIngester()
    assert ingester.can_handle(Path("document.docx"))
    assert not ingester.can_handle(Path("document.doc"))
    assert not ingester.can_handle(Path("document.pdf"))


def test_wraps_paragraphs_in_p_tags(simple_docx, docx_config):
    ingester = DocxIngester()
    raw = ingester.ingest(simple_docx, docx_config)

    assert "<p>" in raw.text
    assert "First paragraph" in raw.text
