"""End-to-end integration tests — run the full pipeline on real files.

These tests exercise Stages 1-6 together. AI is disabled (rewrite_percent=0,
generators off) so no API calls are made.

Per IMPLEMENTATION.md §13.3 success criteria:
  - All input formats → EPUB
  - All output formats pass validation
  - Tables render with borders in DOCX
"""

from __future__ import annotations

import asyncio
from pathlib import Path

import pytest

from bookforge.core.models import BookMetadata, JobConfig
from bookforge.core.pipeline import Pipeline


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def pipeline(tmp_path: Path) -> Pipeline:
    """Pipeline with temp dir and AI disabled."""
    return Pipeline({
        "pipeline": {"temp_dir": str(tmp_path / "temp")},
        "ai": {"provider": "anthropic", "api_key": ""},
    })


@pytest.fixture
def metadata() -> BookMetadata:
    return BookMetadata(
        title="Integration Test Book",
        authors=["Test Author"],
        publisher_name="Test Press",
        year=2026,
        language="en",
    )


@pytest.fixture
def job_config_epub() -> JobConfig:
    return JobConfig(
        output_formats=["epub"],
        generate_title=False,
        generate_preface=False,
        generate_acknowledgement=False,
        rewrite_percent=0,
    )


@pytest.fixture
def job_config_docx() -> JobConfig:
    return JobConfig(
        output_formats=["docx"],
        generate_title=False,
        generate_preface=False,
        generate_acknowledgement=False,
        rewrite_percent=0,
    )


@pytest.fixture
def job_config_multi() -> JobConfig:
    return JobConfig(
        output_formats=["epub", "docx"],
        generate_title=False,
        generate_preface=False,
        generate_acknowledgement=False,
        rewrite_percent=0,
    )


# ---------------------------------------------------------------------------
# Sample file fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def html_file(tmp_path: Path) -> Path:
    content = """<!DOCTYPE html>
<html><head><title>Test</title></head><body>
<h1>Introduction</h1>
<p>This is a test chapter with a table.</p>
<table>
<thead><tr><th>Name</th><th>Value</th></tr></thead>
<tbody><tr><td>Alpha</td><td>1.0</td></tr><tr><td>Beta</td><td>2.0</td></tr></tbody>
</table>
<h2>Conclusion</h2>
<p>Final paragraph.</p>
</body></html>"""
    path = tmp_path / "test.html"
    path.write_text(content)
    return path


@pytest.fixture
def txt_file(tmp_path: Path) -> Path:
    content = """Chapter 1: Introduction

This is the introduction text.

Chapter 2: Methods

This describes the methodology.
"""
    path = tmp_path / "test.txt"
    path.write_text(content)
    return path


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    from docx import Document
    doc = Document()
    doc.add_heading("Test Chapter", level=1)
    doc.add_paragraph("Content of the test chapter.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Val"
    table.rows[1].cells[0].text = "A"
    table.rows[1].cells[1].text = "1"
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Helper
# ---------------------------------------------------------------------------


def run_pipeline(pipeline, input_file, metadata, job_config):
    """Synchronous helper to run the async pipeline."""
    async def _run():
        normalized = await pipeline.process_file(input_file, job_config)
        return await pipeline.process_book([normalized], metadata, job_config)
    return asyncio.run(_run())


# ---------------------------------------------------------------------------
# Tests — HTML input
# ---------------------------------------------------------------------------


def test_html_to_epub(pipeline, html_file, metadata, job_config_epub):
    outputs = run_pipeline(pipeline, html_file, metadata, job_config_epub)
    assert len(outputs) >= 1
    epub_path = outputs[0]
    assert epub_path.suffix == ".epub"
    assert epub_path.exists()
    assert epub_path.stat().st_size > 0


def test_html_to_docx(pipeline, html_file, metadata, job_config_docx):
    outputs = run_pipeline(pipeline, html_file, metadata, job_config_docx)
    assert len(outputs) >= 1
    docx_path = outputs[0]
    assert docx_path.suffix == ".docx"
    assert docx_path.exists()

    # Verify opens with python-docx
    from docx import Document
    doc = Document(str(docx_path))
    assert len(doc.paragraphs) > 0


def test_html_to_multi_format(pipeline, html_file, metadata, job_config_multi):
    outputs = run_pipeline(pipeline, html_file, metadata, job_config_multi)
    suffixes = {p.suffix for p in outputs}
    assert ".epub" in suffixes
    assert ".docx" in suffixes


# ---------------------------------------------------------------------------
# Tests — TXT input
# ---------------------------------------------------------------------------


def test_txt_to_epub(pipeline, txt_file, metadata, job_config_epub):
    outputs = run_pipeline(pipeline, txt_file, metadata, job_config_epub)
    assert len(outputs) >= 1
    assert outputs[0].suffix == ".epub"
    assert outputs[0].stat().st_size > 0


# ---------------------------------------------------------------------------
# Tests — DOCX input
# ---------------------------------------------------------------------------


def test_docx_to_epub(pipeline, docx_file, metadata, job_config_epub):
    outputs = run_pipeline(pipeline, docx_file, metadata, job_config_epub)
    assert len(outputs) >= 1
    assert outputs[0].suffix == ".epub"
    assert outputs[0].stat().st_size > 0


def test_docx_to_docx(pipeline, docx_file, metadata, job_config_docx):
    """DOCX input → DOCX output (re-formatting with template)."""
    outputs = run_pipeline(pipeline, docx_file, metadata, job_config_docx)
    assert len(outputs) >= 1
    assert outputs[0].suffix == ".docx"


# ---------------------------------------------------------------------------
# Tests — DOCX table borders (success criterion §5)
# ---------------------------------------------------------------------------


def test_docx_output_has_table_borders(pipeline, html_file, metadata, job_config_docx):
    """Tables in DOCX output should have hairline grid borders."""
    outputs = run_pipeline(pipeline, html_file, metadata, job_config_docx)
    docx_path = [p for p in outputs if p.suffix == ".docx"][0]

    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(str(docx_path))

    assert len(doc.tables) >= 1, "DOCX should contain at least one table"
    cell = doc.tables[0].rows[0].cells[0]
    borders = cell._tc.findall(".//" + qn("w:tcBorders"))
    assert len(borders) > 0, "Table cell should have border definitions"


# ---------------------------------------------------------------------------
# Tests — multi-file assembly
# ---------------------------------------------------------------------------


def test_multi_file_assembly(pipeline, tmp_path, metadata, job_config_epub):
    """Multiple input files assembled into one book."""
    f1 = tmp_path / "chapter1.html"
    f1.write_text("<h1>Chapter 1</h1><p>First chapter.</p>")
    f2 = tmp_path / "chapter2.html"
    f2.write_text("<h1>Chapter 2</h1><p>Second chapter.</p>")

    async def _run():
        n1 = await pipeline.process_file(f1, job_config_epub)
        n2 = await pipeline.process_file(f2, job_config_epub)
        return await pipeline.process_book([n1, n2], metadata, job_config_epub)

    outputs = asyncio.run(_run())
    assert len(outputs) >= 1
    assert outputs[0].suffix == ".epub"
    assert outputs[0].stat().st_size > 0


# ---------------------------------------------------------------------------
# Tests — EPUB validation
# ---------------------------------------------------------------------------


def test_epub_is_valid_zip(pipeline, html_file, metadata, job_config_epub):
    """EPUB output should be a valid ZIP file."""
    import zipfile
    outputs = run_pipeline(pipeline, html_file, metadata, job_config_epub)
    epub_path = outputs[0]
    assert zipfile.is_zipfile(epub_path)

    with zipfile.ZipFile(epub_path) as zf:
        names = zf.namelist()
        # EPUB must contain certain files
        assert any("content.opf" in n or ".opf" in n for n in names)
